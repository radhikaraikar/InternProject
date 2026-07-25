"""
ingest.py — Document loading, chunking, embedding, and ChromaDB ingestion.

Handles: PDFs (text-native via PyMuPDF), plain text, .docx
OCR: skipped gracefully if pytesseract/pdf2image not installed.
Embeddings: distilbert-base-uncased via local transformers cache (no download).
Vector store: ChromaDB 1.5.x (PersistentClient API).
"""

import logging
import os
import re
from pathlib import Path
from typing import List, Tuple

import numpy as np
import torch
from transformers import AutoModel, AutoTokenizer

import chromadb
import fitz  # PyMuPDF
import docx as _docx

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
logger = logging.getLogger(__name__)

# ── Config ────────────────────────────────────────────────────────────────────
CHROMA_DIR = "./chroma_db"
COLLECTION_NAME = "rag_docs"
CHUNK_SIZE = 800
CHUNK_OVERLAP = 150
OCR_TEXT_THRESHOLD = 30   # chars per page below which we try OCR
EMBED_MODEL = "distilbert-base-uncased"
EMBED_BATCH = 32

# ── Embedding model (loaded once, from local cache) ───────────────────────────
_tokenizer: AutoTokenizer | None = None
_embed_model: AutoModel | None = None


def _load_embed_model():
    global _tokenizer, _embed_model
    if _embed_model is None:
        logger.info(f"Loading embedding model '{EMBED_MODEL}' from local cache…")
        _tokenizer = AutoTokenizer.from_pretrained(EMBED_MODEL, local_files_only=True)
        _embed_model = AutoModel.from_pretrained(EMBED_MODEL, local_files_only=True)
        _embed_model.eval()
        logger.info("Embedding model ready.")
    return _tokenizer, _embed_model


def embed_texts(texts: List[str]) -> List[List[float]]:
    """Attention-masked mean-pool distilbert embeddings (L2-normalised)."""
    tok, model = _load_embed_model()
    all_embeddings: List[List[float]] = []

    for i in range(0, len(texts), EMBED_BATCH):
        batch = texts[i : i + EMBED_BATCH]
        inputs = tok(
            batch,
            padding=True,
            truncation=True,
            max_length=512,
            return_tensors="pt",
        )
        with torch.no_grad():
            out = model(**inputs)

        # Attention-mask weighted mean pool — ignores padding tokens
        mask = inputs["attention_mask"].unsqueeze(-1).float()  # (B, T, 1)
        emb = (out.last_hidden_state * mask).sum(dim=1) / mask.sum(dim=1).clamp(min=1e-9)

        # L2-normalise for cosine similarity
        emb = emb / emb.norm(dim=-1, keepdim=True).clamp(min=1e-9)
        all_embeddings.extend(emb.tolist())

    return all_embeddings


# ── ChromaDB embedding function (wraps embed_texts) ──────────────────────────
class DistilBertEmbeddingFunction(chromadb.EmbeddingFunction):
    def __call__(self, input: chromadb.Documents) -> chromadb.Embeddings:
        return embed_texts(list(input))


def get_embedding_function() -> DistilBertEmbeddingFunction:
    return DistilBertEmbeddingFunction()


# ── Text cleaning ─────────────────────────────────────────────────────────────
def _clean(text: str) -> str:
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"(Page \d+ of \d+|^\d+$)", "", text, flags=re.MULTILINE)
    return text.strip()


# ── PDF extraction (PyMuPDF, with optional OCR fallback) ──────────────────────
def _try_ocr(page: fitz.Page) -> str:
    """Attempt OCR on a page image if pytesseract is available."""
    try:
        import pytesseract
        from PIL import Image
        import io
        pix = page.get_pixmap(dpi=200)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        return pytesseract.image_to_string(img, lang="eng")
    except ImportError:
        return ""
    except Exception as e:
        logger.debug(f"OCR failed: {e}")
        return ""


def extract_pdf(path: str) -> List[Tuple[int, str]]:
    """Return [(page_num, text), ...] for a PDF file."""
    results = []
    doc = fitz.open(path)
    for i, page in enumerate(doc, start=1):
        text = _clean(page.get_text("text") or "")
        if len(text) < OCR_TEXT_THRESHOLD:
            ocr = _clean(_try_ocr(page))
            if len(ocr) > len(text):
                text = ocr
        if text:
            results.append((i, text))
    doc.close()
    return results


def extract_image(path: str) -> str:
    """OCR a standalone image file."""
    try:
        import pytesseract
        from PIL import Image
        return pytesseract.image_to_string(Image.open(path), lang="eng")
    except ImportError:
        logger.warning("pytesseract not installed — skipping image OCR.")
        return ""
    except Exception as e:
        logger.warning(f"Image OCR failed: {e}")
        return ""


def extract_docx(path: str) -> str:
    doc = _docx.Document(path)
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())


def extract_txt(path: str) -> str:
    with open(path, encoding="utf-8", errors="replace") as f:
        return f.read()


# ── Document loading ──────────────────────────────────────────────────────────
def load_documents_from_folder(folder: str) -> List[Document]:
    """Load all supported files from a folder into LangChain Documents."""
    folder = Path(folder)
    supported = {".pdf", ".txt", ".docx", ".png", ".jpg", ".jpeg", ".tiff", ".bmp"}
    files = [f for f in folder.rglob("*") if f.suffix.lower() in supported]

    if not files:
        logger.warning(f"No supported files found in {folder}")
        return []

    docs: List[Document] = []
    for f in files:
        ext = f.suffix.lower()
        logger.info(f"  Loading: {f.name}")
        try:
            if ext == ".pdf":
                for page_num, text in extract_pdf(str(f)):
                    docs.append(Document(page_content=text,
                                         metadata={"source": f.name, "page": page_num}))
            elif ext in {".png", ".jpg", ".jpeg", ".tiff", ".bmp"}:
                text = _clean(extract_image(str(f)))
                if text:
                    docs.append(Document(page_content=text,
                                         metadata={"source": f.name, "page": 1}))
            elif ext == ".docx":
                text = _clean(extract_docx(str(f)))
                if text:
                    docs.append(Document(page_content=text,
                                         metadata={"source": f.name, "page": 1}))
            elif ext == ".txt":
                text = _clean(extract_txt(str(f)))
                if text:
                    docs.append(Document(page_content=text,
                                         metadata={"source": f.name, "page": 1}))
        except Exception as e:
            logger.error(f"  Failed to load {f.name}: {e}")

    logger.info(f"Loaded {len(docs)} pages from {len(files)} files.")
    return docs


# ── Chunking ──────────────────────────────────────────────────────────────────
def chunk_documents(docs: List[Document]) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks = splitter.split_documents(docs)
    for i, c in enumerate(chunks):
        c.metadata["chunk_id"] = f"chunk_{i:05d}"
    logger.info(f"Chunked into {len(chunks)} pieces.")
    return chunks


# ── ChromaDB upsert ───────────────────────────────────────────────────────────
def ingest_to_chroma(chunks: List[Document], persist_dir: str = CHROMA_DIR) -> chromadb.Collection:
    ef = get_embedding_function()
    client = chromadb.PersistentClient(path=persist_dir)

    # Drop and recreate for clean re-ingestion
    try:
        client.delete_collection(COLLECTION_NAME)
        logger.info("Dropped existing collection.")
    except Exception:
        pass

    col = client.create_collection(
        name=COLLECTION_NAME,
        embedding_function=ef,
        metadata={"hnsw:space": "cosine"},
    )

    batch = 64
    for i in range(0, len(chunks), batch):
        b = chunks[i : i + batch]
        col.upsert(
            ids=[c.metadata["chunk_id"] for c in b],
            documents=[c.page_content for c in b],
            metadatas=[c.metadata for c in b],
        )
        logger.info(f"  Upserted {min(i + batch, len(chunks))}/{len(chunks)} chunks…")

    logger.info(f"✅ Ingestion complete — {col.count()} chunks in ChromaDB.")
    return col


# ── Public helpers ────────────────────────────────────────────────────────────
def run_ingestion(data_folder: str = "./data", persist_dir: str = CHROMA_DIR) -> chromadb.Collection | None:
    logger.info(f"Ingesting '{data_folder}' → '{persist_dir}'")
    docs = load_documents_from_folder(data_folder)
    if not docs:
        logger.error("No documents loaded.")
        return None
    chunks = chunk_documents(docs)
    return ingest_to_chroma(chunks, persist_dir)


def get_chroma_collection(persist_dir: str = CHROMA_DIR) -> chromadb.Collection:
    ef = get_embedding_function()
    client = chromadb.PersistentClient(path=persist_dir)
    return client.get_collection(name=COLLECTION_NAME, embedding_function=ef)


if __name__ == "__main__":
    run_ingestion()
